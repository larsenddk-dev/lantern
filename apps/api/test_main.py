"""
Lantern API tests — no network, no secrets required.

The OpenAI-compatible streaming call is monkeypatched so tests run fully
offline. Tests prove:

1. /health returns {"status": "ok", ...}
2. Session round-trip: create → list → get with messages.
3. POST /chat streams reply in multiple SSE chunks (using a fake streaming
   source), persists both user and assistant messages.
4. Provider CRUD: create → list → update → delete.
5. Active provider selection persists across reads.
6. /chat uses the selected provider's base_url + model (stub verifies routing).
7. API keys are never returned in full (masked on read).
"""

import json
import os
import tempfile
import pytest

# Point the DB at a fresh temp file before importing the app so no production
# data is touched.
_tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
_tmp.close()
os.environ["LANTERN_DB_PATH"] = _tmp.name

# Minimal provider env — not used in tests (client is monkeypatched)
os.environ.setdefault("LANTERN_OPENAI_API_KEY", "test-key")
os.environ.setdefault("LANTERN_MODEL", "test-model")

import main  # noqa: E402 — imported after env patching
from fastapi.testclient import TestClient  # noqa: E402

# Call init_db explicitly so TestClient requests hit an initialised schema.
main.init_db()

client = TestClient(main.app, raise_server_exceptions=True)

# ---------------------------------------------------------------------------
# Fake streaming source — yields three deltas
# ---------------------------------------------------------------------------

FAKE_DELTAS = ["Hello", ", ", "world!"]

# Tracks what the last stream_chat_completion call received
_last_stream_kwargs: dict = {}


async def _fake_stream(*_args, **kwargs):
    _last_stream_kwargs.clear()
    _last_stream_kwargs.update(kwargs)
    for token in FAKE_DELTAS:
        yield token


# ---------------------------------------------------------------------------
# Tests — original Phase 1
# ---------------------------------------------------------------------------


def test_health():
    resp = client.get("/health")
    assert resp.status_code == 200
    body = resp.json()
    assert body["status"] == "ok"


def test_create_and_list_sessions():
    resp = client.post("/sessions", json={"title": "Test session"})
    assert resp.status_code == 200
    data = resp.json()
    assert data["title"] == "Test session"
    assert "id" in data

    list_resp = client.get("/sessions")
    assert list_resp.status_code == 200
    ids = [s["id"] for s in list_resp.json()]
    assert data["id"] in ids


def test_get_session_not_found():
    resp = client.get("/sessions/nonexistent-id")
    assert resp.status_code == 404


def test_rename_session():
    sid = client.post("/sessions", json={"title": "Before"}).json()["id"]
    resp = client.patch(f"/sessions/{sid}", json={"title": "After"})
    assert resp.status_code == 200
    assert resp.json()["title"] == "After"
    listed = client.get("/sessions").json()
    assert any(s["id"] == sid and s["title"] == "After" for s in listed)


def test_rename_session_not_found():
    resp = client.patch("/sessions/nope", json={"title": "x"})
    assert resp.status_code == 404


def test_delete_session():
    sid = client.post("/sessions", json={"title": "To delete"}).json()["id"]
    resp = client.delete(f"/sessions/{sid}")
    assert resp.status_code == 200
    assert resp.json()["ok"] is True
    assert client.get(f"/sessions/{sid}").status_code == 404


def test_delete_session_not_found():
    resp = client.delete("/sessions/nope")
    assert resp.status_code == 404


def test_chat_streaming_and_session_history(monkeypatch):
    """End-to-end: create session, send message, verify streaming chunks,
    then reload session and confirm both messages are stored."""
    monkeypatch.setattr(main, "stream_chat_completion", _fake_stream)

    # Create session
    sess = client.post("/sessions", json={"title": "Streaming test"}).json()
    session_id = sess["id"]

    # Send a chat message — stream response
    with client.stream(
        "POST",
        "/chat",
        json={"session_id": session_id, "message": "Say hello"},
    ) as resp:
        assert resp.status_code == 200
        assert "text/event-stream" in resp.headers["content-type"]

        raw = b"".join(resp.iter_bytes())

    lines = raw.decode().strip().splitlines()

    # Parse SSE lines: "data: <json>" and "data: [DONE]"
    data_lines = [l[len("data: "):] for l in lines if l.startswith("data: ")]
    delta_lines = [l for l in data_lines if l != "[DONE]"]
    done_lines = [l for l in data_lines if l == "[DONE]"]

    assert len(done_lines) == 1, "Expected exactly one [DONE] sentinel"
    assert len(delta_lines) >= 2, (
        f"Expected multiple delta chunks, got {len(delta_lines)}: {delta_lines}"
    )

    deltas = [json.loads(l)["delta"] for l in delta_lines]
    assert "".join(deltas) == "Hello, world!"

    # Reload session and verify history
    reloaded = client.get(f"/sessions/{session_id}").json()
    messages = reloaded["messages"]
    assert len(messages) == 2, f"Expected 2 messages, got {len(messages)}"
    assert messages[0]["role"] == "user"
    assert messages[0]["content"] == "Say hello"
    assert messages[1]["role"] == "assistant"
    assert messages[1]["content"] == "Hello, world!"


def test_chat_unknown_session():
    resp = client.post(
        "/chat", json={"session_id": "no-such-session", "message": "hi"}
    )
    assert resp.status_code == 404


# ---------------------------------------------------------------------------
# Tests — Phase 2a: provider CRUD
# ---------------------------------------------------------------------------


def test_provider_create_list_update_delete():
    """CRUD round-trip for provider configs."""
    # Create
    resp = client.post("/providers", json={
        "label": "My OpenRouter",
        "base_url": "https://openrouter.ai/api/v1",
        "model": "openai/gpt-4o-mini",
        "api_key": "sk-secret-key",
    })
    assert resp.status_code == 200
    provider = resp.json()
    pid = provider["id"]
    assert provider["label"] == "My OpenRouter"
    assert provider["base_url"] == "https://openrouter.ai/api/v1"
    assert provider["model"] == "openai/gpt-4o-mini"
    # Key must be masked — never the real value
    assert "sk-secret-key" not in json.dumps(provider)
    assert provider["api_key_masked"] == "sk-***"
    assert provider["is_active"] is False

    # List — should contain our new provider
    list_resp = client.get("/providers")
    assert list_resp.status_code == 200
    ids = [p["id"] for p in list_resp.json()]
    assert pid in ids

    # Update label and model
    upd_resp = client.put(f"/providers/{pid}", json={
        "label": "OpenRouter Free",
        "model": "openai/gpt-4o",
    })
    assert upd_resp.status_code == 200
    updated = upd_resp.json()
    assert updated["label"] == "OpenRouter Free"
    assert updated["model"] == "openai/gpt-4o"
    # base_url unchanged
    assert updated["base_url"] == "https://openrouter.ai/api/v1"

    # Delete
    del_resp = client.delete(f"/providers/{pid}")
    assert del_resp.status_code == 200
    assert del_resp.json()["ok"] is True

    # Confirm gone
    list_after = [p["id"] for p in client.get("/providers").json()]
    assert pid not in list_after


def test_provider_not_found():
    resp = client.put("/providers/no-such-id", json={"label": "x"})
    assert resp.status_code == 404

    resp2 = client.delete("/providers/no-such-id")
    assert resp2.status_code == 404

    resp3 = client.post("/providers/no-such-id/activate")
    assert resp3.status_code == 404


# ---------------------------------------------------------------------------
# Tests — Phase 2a: active provider selection + persistence
# ---------------------------------------------------------------------------


def test_active_provider_selection_persists():
    """Create two providers, activate one, confirm active selection persists."""
    # Create provider A
    pa = client.post("/providers", json={
        "label": "Provider A",
        "base_url": "https://a.example.com/v1",
        "model": "model-a",
        "api_key": "key-a",
    }).json()

    # Create provider B
    pb = client.post("/providers", json={
        "label": "Provider B",
        "base_url": "https://b.example.com/v1",
        "model": "model-b",
        "api_key": "key-b",
    }).json()

    # Activate A
    act_resp = client.post(f"/providers/{pa['id']}/activate")
    assert act_resp.status_code == 200
    assert act_resp.json()["is_active"] is True

    # GET /providers/active should return A
    active_resp = client.get("/providers/active")
    assert active_resp.status_code == 200
    active = active_resp.json()["active"]
    assert active is not None
    assert active["id"] == pa["id"]
    assert active["label"] == "Provider A"
    # Key must not be leaked
    assert "key-a" not in json.dumps(active)

    # Activate B — A should become inactive
    client.post(f"/providers/{pb['id']}/activate")
    active_resp2 = client.get("/providers/active")
    active2 = active_resp2.json()["active"]
    assert active2["id"] == pb["id"]

    # List should show only B as active
    providers = client.get("/providers").json()
    pa_row = next(p for p in providers if p["id"] == pa["id"])
    pb_row = next(p for p in providers if p["id"] == pb["id"])
    assert pa_row["is_active"] is False
    assert pb_row["is_active"] is True

    # Clean up
    client.delete(f"/providers/{pa['id']}")
    client.delete(f"/providers/{pb['id']}")


# ---------------------------------------------------------------------------
# Tests — Phase 2a: /chat uses selected provider/model (stub verifies routing)
# ---------------------------------------------------------------------------


def test_chat_uses_active_provider(monkeypatch):
    """Verify that /chat routes to the active provider's base_url and model."""
    monkeypatch.setattr(main, "stream_chat_completion", _fake_stream)

    # Create a provider and activate it
    prov = client.post("/providers", json={
        "label": "Stub Provider",
        "base_url": "https://stub.example.com/v1",
        "model": "stub-model-1",
        "api_key": "stub-key",
    }).json()
    client.post(f"/providers/{prov['id']}/activate")

    # Create a session and send a chat
    sess = client.post("/sessions", json={"title": "Provider routing test"}).json()
    with client.stream("POST", "/chat", json={
        "session_id": sess["id"],
        "message": "ping",
    }) as resp:
        assert resp.status_code == 200
        b"".join(resp.iter_bytes())  # consume stream

    # Confirm the stub was called with the active provider's base_url + model
    assert _last_stream_kwargs.get("base_url") == "https://stub.example.com/v1"
    assert _last_stream_kwargs.get("model") == "stub-model-1"

    # Clean up
    client.delete(f"/providers/{prov['id']}")


def test_chat_uses_env_fallback_when_no_active_provider(monkeypatch):
    """When no provider is active, /chat falls back to env vars."""
    monkeypatch.setattr(main, "stream_chat_completion", _fake_stream)

    # Ensure no active provider (deactivate all by activating a throwaway then deleting it)
    tmp = client.post("/providers", json={
        "label": "Throwaway",
        "base_url": "https://throwaway.example.com/v1",
        "model": "throwaway-model",
        "api_key": "",
    }).json()
    client.post(f"/providers/{tmp['id']}/activate")
    client.delete(f"/providers/{tmp['id']}")

    # Confirm no active provider in DB
    active_resp = client.get("/providers/active")
    assert active_resp.json()["active"] is None

    sess = client.post("/sessions", json={"title": "Env fallback test"}).json()
    with client.stream("POST", "/chat", json={
        "session_id": sess["id"],
        "message": "fallback ping",
    }) as resp:
        assert resp.status_code == 200
        b"".join(resp.iter_bytes())

    # Should fall back to the env-configured values set at top of this file
    assert _last_stream_kwargs.get("base_url") == main.LANTERN_OPENAI_BASE_URL
    assert _last_stream_kwargs.get("model") == main.LANTERN_MODEL


# ---------------------------------------------------------------------------
# Tests — Phase 2a: API key masking
# ---------------------------------------------------------------------------


def test_api_key_never_returned_in_full():
    """Keys stored must never appear in API responses."""
    real_key = "super-secret-api-key-12345"
    prov = client.post("/providers", json={
        "label": "Key Masking Test",
        "base_url": "https://example.com/v1",
        "model": "test-model",
        "api_key": real_key,
    }).json()
    pid = prov["id"]

    # Single provider response — key must not appear
    assert real_key not in json.dumps(prov)

    # List response — key must not appear
    providers_json = json.dumps(client.get("/providers").json())
    assert real_key not in providers_json

    # Update response — key must not appear even after updating
    upd = client.put(f"/providers/{pid}", json={"label": "Updated"}).json()
    assert real_key not in json.dumps(upd)

    # Activate response — key must not appear
    act = client.post(f"/providers/{pid}/activate").json()
    assert real_key not in json.dumps(act)

    # Active provider response — key must not appear
    active_json = json.dumps(client.get("/providers/active").json())
    assert real_key not in active_json

    # Clean up
    client.delete(f"/providers/{pid}")


# ---------------------------------------------------------------------------
# Tests — Phase 2b: Notes CRUD
# ---------------------------------------------------------------------------


def test_notes_create_list_get_update_delete():
    """Full CRUD round-trip for notes."""
    # Create
    resp = client.post("/notes", json={"title": "My First Note", "content": "Hello world"})
    assert resp.status_code == 200
    note = resp.json()
    nid = note["id"]
    assert note["title"] == "My First Note"
    assert note["content"] == "Hello world"
    assert "created_at" in note
    assert "updated_at" in note

    # List — should contain our new note
    list_resp = client.get("/notes")
    assert list_resp.status_code == 200
    ids = [n["id"] for n in list_resp.json()]
    assert nid in ids

    # Get single note
    get_resp = client.get(f"/notes/{nid}")
    assert get_resp.status_code == 200
    fetched = get_resp.json()
    assert fetched["id"] == nid
    assert fetched["title"] == "My First Note"
    assert fetched["content"] == "Hello world"

    # Update title only (partial update)
    upd_resp = client.put(f"/notes/{nid}", json={"title": "Updated Title"})
    assert upd_resp.status_code == 200
    updated = upd_resp.json()
    assert updated["title"] == "Updated Title"
    assert updated["content"] == "Hello world"  # unchanged

    # Update content only
    upd2_resp = client.put(f"/notes/{nid}", json={"content": "New content"})
    assert upd2_resp.status_code == 200
    updated2 = upd2_resp.json()
    assert updated2["title"] == "Updated Title"
    assert updated2["content"] == "New content"

    # Delete
    del_resp = client.delete(f"/notes/{nid}")
    assert del_resp.status_code == 200
    assert del_resp.json()["ok"] is True

    # Confirm gone
    get_after = client.get(f"/notes/{nid}")
    assert get_after.status_code == 404

    ids_after = [n["id"] for n in client.get("/notes").json()]
    assert nid not in ids_after


def test_notes_not_found():
    """404 on get/update/delete of non-existent note."""
    resp = client.get("/notes/no-such-id")
    assert resp.status_code == 404

    resp2 = client.put("/notes/no-such-id", json={"title": "x"})
    assert resp2.status_code == 404

    resp3 = client.delete("/notes/no-such-id")
    assert resp3.status_code == 404


def test_notes_empty_fields_allowed():
    """Notes can be created with empty title and content (blank note)."""
    resp = client.post("/notes", json={})
    assert resp.status_code == 200
    note = resp.json()
    assert note["title"] == ""
    assert note["content"] == ""
    # Clean up
    client.delete(f"/notes/{note['id']}")


# ---------------------------------------------------------------------------
# Tests — Phase 2b: Tasks CRUD + done toggle
# ---------------------------------------------------------------------------


def test_tasks_create_list_toggle_delete():
    """Full CRUD round-trip for tasks including done/undone toggle."""
    # Create
    resp = client.post("/tasks", json={"title": "Buy groceries"})
    assert resp.status_code == 200
    task = resp.json()
    tid = task["id"]
    assert task["title"] == "Buy groceries"
    assert task["done"] is False
    assert "created_at" in task
    assert "updated_at" in task

    # List — should contain our new task
    list_resp = client.get("/tasks")
    assert list_resp.status_code == 200
    ids = [t["id"] for t in list_resp.json()]
    assert tid in ids

    # Get single task
    get_resp = client.get(f"/tasks/{tid}")
    assert get_resp.status_code == 200
    fetched = get_resp.json()
    assert fetched["id"] == tid
    assert fetched["done"] is False

    # Toggle done = True
    toggle_resp = client.put(f"/tasks/{tid}", json={"done": True})
    assert toggle_resp.status_code == 200
    toggled = toggle_resp.json()
    assert toggled["done"] is True
    assert toggled["title"] == "Buy groceries"  # unchanged

    # Toggle done = False (undone)
    toggle2_resp = client.put(f"/tasks/{tid}", json={"done": False})
    assert toggle2_resp.status_code == 200
    toggled2 = toggle2_resp.json()
    assert toggled2["done"] is False

    # Update title
    upd_resp = client.put(f"/tasks/{tid}", json={"title": "Buy organic groceries"})
    assert upd_resp.status_code == 200
    updated = upd_resp.json()
    assert updated["title"] == "Buy organic groceries"
    assert updated["done"] is False  # unchanged

    # Delete
    del_resp = client.delete(f"/tasks/{tid}")
    assert del_resp.status_code == 200
    assert del_resp.json()["ok"] is True

    # Confirm gone
    get_after = client.get(f"/tasks/{tid}")
    assert get_after.status_code == 404

    ids_after = [t["id"] for t in client.get("/tasks").json()]
    assert tid not in ids_after


def test_tasks_not_found():
    """404 on get/update/delete of non-existent task."""
    resp = client.get("/tasks/no-such-id")
    assert resp.status_code == 404

    resp2 = client.put("/tasks/no-such-id", json={"done": True})
    assert resp2.status_code == 404

    resp3 = client.delete("/tasks/no-such-id")
    assert resp3.status_code == 404


def test_tasks_done_field_is_boolean():
    """Ensure done field is always serialized as boolean, not integer."""
    resp = client.post("/tasks", json={"title": "Type-check task"})
    assert resp.status_code == 200
    task = resp.json()
    tid = task["id"]
    assert task["done"] is False  # must be bool False, not 0

    toggled = client.put(f"/tasks/{tid}", json={"done": True}).json()
    assert toggled["done"] is True  # must be bool True, not 1

    # Verify via list too
    tasks_list = client.get("/tasks").json()
    matching = [t for t in tasks_list if t["id"] == tid]
    assert len(matching) == 1
    assert matching[0]["done"] is True

    # Clean up
    client.delete(f"/tasks/{tid}")


# ---------------------------------------------------------------------------
# Tests — Phase 2c: Documents (upload + text extraction)
# ---------------------------------------------------------------------------


def test_documents_upload_list_get_delete_text():
    """Upload a .txt, list it, fetch extracted text, then delete it."""
    content = b"Hello from a Lantern text file.\nSecond line."
    resp = client.post("/documents", files={"file": ("note.txt", content, "text/plain")})
    assert resp.status_code == 200
    doc = resp.json()
    assert doc["filename"] == "note.txt"
    assert doc["size_bytes"] == len(content)
    assert doc["has_text"] is True
    assert "Lantern text file" in doc["extracted_text"]
    doc_id = doc["id"]

    # List returns metadata only (no full text field)
    listed = client.get("/documents").json()
    match = [d for d in listed if d["id"] == doc_id]
    assert len(match) == 1
    assert "extracted_text" not in match[0]
    assert match[0]["has_text"] is True

    # Get returns the extracted text
    got = client.get(f"/documents/{doc_id}").json()
    assert "Second line." in got["extracted_text"]

    # Delete, then 404
    assert client.delete(f"/documents/{doc_id}").status_code == 200
    assert client.get(f"/documents/{doc_id}").status_code == 404


def test_documents_markdown_extraction():
    md = b"# Heading\n\nSome **markdown** body text."
    doc = client.post("/documents", files={"file": ("readme.md", md, "text/markdown")}).json()
    assert doc["has_text"] is True
    assert "Heading" in doc["extracted_text"]
    client.delete(f"/documents/{doc['id']}")


def test_documents_docx_extraction():
    """python-docx round-trip: build a .docx in memory, upload, extract its text."""
    import io
    import docx

    buf = io.BytesIO()
    document = docx.Document()
    document.add_paragraph("Hello from a docx paragraph.")
    document.add_paragraph("Lantern second paragraph.")
    document.save(buf)
    buf.seek(0)

    doc = client.post(
        "/documents",
        files={"file": (
            "report.docx",
            buf.read(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )},
    ).json()
    assert doc["has_text"] is True
    assert "Hello from a docx paragraph." in doc["extracted_text"]
    assert "Lantern second paragraph." in doc["extracted_text"]
    client.delete(f"/documents/{doc['id']}")


def test_documents_pdf_upload_is_robust():
    """A .pdf uploads/lists/deletes without crashing even if extraction yields nothing."""
    pdf_bytes = b"%PDF-1.4\n1 0 obj<</Type/Catalog>>endobj\ntrailer<</Root 1 0 R>>\n%%EOF"
    resp = client.post("/documents", files={"file": ("doc.pdf", pdf_bytes, "application/pdf")})
    assert resp.status_code == 200
    doc = resp.json()
    assert doc["filename"] == "doc.pdf"
    assert isinstance(doc["has_text"], bool)  # extraction never raises, always a bool
    assert client.delete(f"/documents/{doc['id']}").status_code == 200


def test_documents_filename_is_sanitized():
    """Path-traversal filenames are reduced to a safe basename."""
    doc = client.post(
        "/documents",
        files={"file": ("../../etc/evil.txt", b"x", "text/plain")},
    ).json()
    assert "/" not in doc["filename"]
    assert ".." not in doc["filename"]
    assert doc["filename"] == "evil.txt"
    client.delete(f"/documents/{doc['id']}")


def test_documents_not_found():
    assert client.get("/documents/nonexistent-id").status_code == 404
    assert client.delete("/documents/nonexistent-id").status_code == 404


def test_documents_empty_file_rejected():
    resp = client.post("/documents", files={"file": ("empty.txt", b"", "text/plain")})
    assert resp.status_code == 400


# ---------------------------------------------------------------------------
# Tests — Memory module
# ---------------------------------------------------------------------------


def test_memories_create_list_get_update_delete():
    created = client.post("/memories", json={"content": "User prefers dark mode."}).json()
    assert created["content"] == "User prefers dark mode."
    assert created["pinned"] is False
    mid = created["id"]

    # Get
    got = client.get(f"/memories/{mid}").json()
    assert got["content"] == "User prefers dark mode."

    # Update content + pin
    upd = client.put(f"/memories/{mid}", json={"content": "User prefers light mode.", "pinned": True}).json()
    assert upd["content"] == "User prefers light mode."
    assert upd["pinned"] is True

    # List contains it
    listed = client.get("/memories").json()
    assert any(m["id"] == mid for m in listed)

    # Delete → 404
    assert client.delete(f"/memories/{mid}").status_code == 200
    assert client.get(f"/memories/{mid}").status_code == 404


def test_memories_pinned_sorted_first():
    a = client.post("/memories", json={"content": "unpinned A"}).json()
    b = client.post("/memories", json={"content": "pinned B", "pinned": True}).json()
    listed = client.get("/memories").json()
    ids = [m["id"] for m in listed]
    # Pinned item must come before the unpinned one
    assert ids.index(b["id"]) < ids.index(a["id"])
    client.delete(f"/memories/{a['id']}")
    client.delete(f"/memories/{b['id']}")


def test_memories_not_found():
    assert client.get("/memories/nope").status_code == 404
    assert client.put("/memories/nope", json={"content": "x"}).status_code == 404
    assert client.delete("/memories/nope").status_code == 404


# ---------------------------------------------------------------------------
# Tests — RAG (embeddings + retrieval), offline with a deterministic stub
# ---------------------------------------------------------------------------


def _stub_embed(texts, **_kwargs):
    """Deterministic bag-of-words hashing embedder — no network, no secrets.
    Shared words → overlapping dimensions → higher cosine similarity."""
    import hashlib

    dim = 96
    vectors = []
    for t in texts:
        v = [0.0] * dim
        for tok in t.lower().split():
            h = int(hashlib.md5(tok.encode()).hexdigest(), 16)
            v[h % dim] += 1.0
        vectors.append(v)
    return vectors


def test_rag_index_and_search(monkeypatch):
    monkeypatch.setattr(main, "embed_texts", _stub_embed)
    m1 = client.post("/memories", json={"content": "The user's favorite programming language is Rust"}).json()
    m2 = client.post("/memories", json={"content": "The user has a golden retriever named Biscuit"}).json()

    idx = client.post("/rag/index").json()
    assert idx["total"] >= 2

    results = client.post("/rag/search", json={"query": "favorite programming language", "k": 3}).json()
    assert len(results) >= 1
    assert "Rust" in results[0]["content"]  # most relevant chunk ranks first

    client.delete(f"/memories/{m1['id']}")
    client.delete(f"/memories/{m2['id']}")


def test_chat_injects_pinned_memory_context(monkeypatch):
    captured = {}

    async def _capture_stream(messages, **_kwargs):
        captured["messages"] = messages
        yield "ok"

    monkeypatch.setattr(main, "stream_chat_completion", _capture_stream)
    monkeypatch.setattr(main, "embed_texts", _stub_embed)

    mem = client.post("/memories", json={"content": "ALWAYS reply in French", "pinned": True}).json()
    sess = client.post("/sessions", json={"title": "ctx"}).json()
    with client.stream("POST", "/chat", json={"session_id": sess["id"], "message": "hello"}) as resp:
        assert resp.status_code == 200
        b"".join(resp.iter_bytes())

    msgs = captured["messages"]
    assert any(m["role"] == "system" and "French" in m["content"] for m in msgs)
    client.delete(f"/memories/{mem['id']}")


def test_chat_use_context_false_skips_injection(monkeypatch):
    captured = {}

    async def _capture_stream(messages, **_kwargs):
        captured["messages"] = messages
        yield "ok"

    monkeypatch.setattr(main, "stream_chat_completion", _capture_stream)
    mem = client.post("/memories", json={"content": "SECRET PHRASE xyzzy", "pinned": True}).json()
    sess = client.post("/sessions", json={"title": "noctx"}).json()
    with client.stream("POST", "/chat",
                       json={"session_id": sess["id"], "message": "hi", "use_context": False}) as resp:
        b"".join(resp.iter_bytes())

    msgs = captured["messages"]
    assert not any("xyzzy" in m["content"] for m in msgs)
    client.delete(f"/memories/{mem['id']}")


# ---------------------------------------------------------------------------
# Tests — Compare (multi-model), offline stub
# ---------------------------------------------------------------------------


def test_compare_returns_one_result_per_target(monkeypatch):
    def _stub_complete(messages, **kwargs):
        return f"reply from {kwargs.get('model')}"

    monkeypatch.setattr(main, "complete_chat_once", _stub_complete)
    p1 = client.post("/providers", json={
        "label": "A", "base_url": "http://a/v1", "model": "model-a", "api_key": "x"}).json()
    p2 = client.post("/providers", json={
        "label": "B", "base_url": "http://b/v1", "model": "model-b", "api_key": "y"}).json()

    resp = client.post("/compare", json={
        "message": "hello",
        "targets": [{"provider_id": p1["id"]}, {"provider_id": p2["id"]}],
    }).json()
    assert len(resp["results"]) == 2
    assert "model-a" in resp["results"][0]["reply"]
    assert "model-b" in resp["results"][1]["reply"]

    client.delete(f"/providers/{p1['id']}")
    client.delete(f"/providers/{p2['id']}")


def test_compare_captures_per_target_errors(monkeypatch):
    def _boom(messages, **kwargs):
        raise RuntimeError("provider down")

    monkeypatch.setattr(main, "complete_chat_once", _boom)
    resp = client.post("/compare", json={"message": "hi", "targets": [{"model": "x"}]}).json()
    assert len(resp["results"]) == 1
    assert resp["results"][0]["error"] is not None
    assert resp["results"][0]["reply"] == ""


# ---------------------------------------------------------------------------
# Tests — Agent (tool-calling loop), offline stub
# ---------------------------------------------------------------------------


def test_agent_calls_tool_then_answers(monkeypatch):
    calls = {"n": 0}

    def _stub_tools(messages, tools, **_kwargs):
        calls["n"] += 1
        if calls["n"] == 1:
            return {
                "role": "assistant",
                "content": None,
                "tool_calls": [{
                    "id": "call_1",
                    "type": "function",
                    "function": {"name": "calculator", "arguments": '{"expression": "2 + 3 * 4"}'},
                }],
            }
        return {"role": "assistant", "content": "The answer is 14."}

    monkeypatch.setattr(main, "chat_with_tools", _stub_tools)
    resp = client.post("/agent", json={"message": "what is 2 + 3 * 4?"}).json()
    assert "14" in resp["reply"]
    assert any(s["tool"] == "calculator" for s in resp["steps"])
    assert resp["steps"][0]["result"] == "14"


def test_agent_answers_without_tools(monkeypatch):
    def _stub_tools(messages, tools, **_kwargs):
        return {"role": "assistant", "content": "Hello there."}

    monkeypatch.setattr(main, "chat_with_tools", _stub_tools)
    resp = client.post("/agent", json={"message": "hi"}).json()
    assert resp["reply"] == "Hello there."
    assert resp["steps"] == []


def test_safe_eval_rejects_non_arithmetic():
    import pytest as _pytest
    assert main._safe_eval("2 + 3 * 4") == 14
    assert main._safe_eval("-(5) + 2 ** 3") == 3
    with _pytest.raises(Exception):
        main._safe_eval("__import__('os').system('echo hi')")
    with _pytest.raises(Exception):
        main._safe_eval("open('/etc/passwd')")


# ---------------------------------------------------------------------------
# Tests — Deep Research (plan → gather → synthesize), offline stub
# ---------------------------------------------------------------------------


def test_research_plans_gathers_and_synthesizes(monkeypatch):
    calls = {"n": 0}

    def _stub_complete(messages, **_kwargs):
        calls["n"] += 1
        if calls["n"] == 1:
            return '["What is X?", "How does X work?"]'  # plan
        return "# Report\n\nA structured report about X."  # synthesis

    monkeypatch.setattr(main, "complete_chat_once", _stub_complete)
    monkeypatch.setattr(main, "embed_texts", _stub_embed)

    resp = client.post("/research", json={"question": "Tell me about X"}).json()
    assert resp["subquestions"] == ["What is X?", "How does X work?"]
    assert "structured report about X" in resp["report"]
    assert len(resp["findings"]) == 2
    assert calls["n"] == 2  # exactly one plan + one synthesis call


def test_parse_subquestions_fallback_for_non_json():
    subs = main._parse_subquestions("1. First sub question here\n2. Second sub question here", 4)
    assert len(subs) == 2
    assert subs[0].startswith("First sub question")


# ---------------------------------------------------------------------------
# Tests — Global search
# ---------------------------------------------------------------------------


def test_search_across_modules():
    n = client.post("/notes", json={"title": "Zebra plan", "content": "about zebras"}).json()
    t = client.post("/tasks", json={"title": "buy zebra food"}).json()
    m = client.post("/memories", json={"content": "user likes zebras"}).json()

    results = client.get("/search", params={"q": "zebra"}).json()["results"]
    types = {r["type"] for r in results}
    assert "note" in types
    assert "task" in types
    assert "memory" in types
    assert all("path" in r and "title" in r for r in results)

    client.delete(f"/notes/{n['id']}")
    client.delete(f"/tasks/{t['id']}")
    client.delete(f"/memories/{m['id']}")


def test_search_empty_query_returns_empty():
    assert client.get("/search", params={"q": "   "}).json()["results"] == []


# ---------------------------------------------------------------------------
# Tests — Web search (env-keyed; not configured in tests)
# ---------------------------------------------------------------------------


def test_web_search_not_configured_by_default():
    # No LANTERN_TAVILY_API_KEY in the test env → graceful, never raises.
    r = main.web_search("anything")
    assert r["configured"] is False
    assert r["results"] == []
    assert "note" in r


def test_agent_web_search_tool_reports_unconfigured(monkeypatch):
    calls = {"n": 0}

    def _stub_tools(messages, tools, **_kwargs):
        calls["n"] += 1
        if calls["n"] == 1:
            return {
                "role": "assistant",
                "content": None,
                "tool_calls": [{
                    "id": "c1", "type": "function",
                    "function": {"name": "web_search", "arguments": '{"query": "latest news"}'},
                }],
            }
        return {"role": "assistant", "content": "I can't search the web yet."}

    monkeypatch.setattr(main, "chat_with_tools", _stub_tools)
    resp = client.post("/agent", json={"message": "what's the news?"}).json()
    assert any(s["tool"] == "web_search" for s in resp["steps"])
    assert "not configured" in resp["steps"][0]["result"].lower()


def test_agent_web_search_with_stubbed_results(monkeypatch):
    def _stub_web(query, max_results=5):
        return {"configured": True, "results": [
            {"title": "Result A", "url": "https://a.example", "content": "alpha content"},
        ]}

    calls = {"n": 0}

    def _stub_tools(messages, tools, **_kwargs):
        calls["n"] += 1
        if calls["n"] == 1:
            return {
                "role": "assistant", "content": None,
                "tool_calls": [{"id": "c1", "type": "function",
                                "function": {"name": "web_search", "arguments": '{"query": "alpha"}'}}],
            }
        return {"role": "assistant", "content": "Found Result A."}

    monkeypatch.setattr(main, "web_search", _stub_web)
    monkeypatch.setattr(main, "chat_with_tools", _stub_tools)
    resp = client.post("/agent", json={"message": "search alpha"}).json()
    assert "Result A" in resp["steps"][0]["result"]
